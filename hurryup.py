#!/usr/bin/env python3

import os
import subprocess
from docx import Document
import getpass
import time

def check_sudo():
    if os.geteuid() != 0:
        print("This tool requires superuser privileges. Please run with sudo.")
        exit(1)

def print_banner():
    print("\n" + "="*50)
    print("SANKET DEORE presents: 'HurryUp'")
    print("Community: Coffee-while-surfing")
    print("="*50 + "\n")
    print("  🐼  Salute! Let's get to work!")
    print("\n")

def list_tools():
    print("Tools available:")
    with open('tools/tool_list.txt', 'r') as file:
        tools = file.readlines()
        for i, tool in enumerate(tools):
            print(f"[{i + 1}] {tool.strip()}")
    return tools

def get_tool_choice(tools):
    choice = int(input("\nEnter the Command no. you want to run.... sweety ;) : "))
    return tools[choice - 1].strip()

def run_tool(tool):
    print(f"\nRunning {tool}...\n")
    result = subprocess.run([tool, '--help'], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    return result.stdout + result.stderr

def create_report(tool, output, format_choice):
    filename = f"reports/{tool}_report.{format_choice}"
    os.makedirs('reports', exist_ok=True)
    
    if format_choice == "docx":
        doc = Document()
        doc.add_heading(f"{tool} Command Report", 0)
        doc.add_paragraph(output)
        doc.save(filename)
    elif format_choice == "pdf":
        # For PDF creation, a third-party library like ReportLab or pypandoc can be used.
        pass

    return os.path.abspath(filename)

def main():
    check_sudo()
    print_banner()
    
    tools = list_tools()
    tool_choice = get_tool_choice(tools)
    output = run_tool(tool_choice)
    
    format_choice = input("\nYour operation is completed successfully. Please select the report file format:\n\n[1] Word\n[2] PDF\n\nEnter the Pattern no: ").strip()
    format_choice = "docx" if format_choice == "1" else "pdf"
    
    report_file = create_report(tool_choice, output, format_choice)
    print(f"\nReport generated: {report_file}")

if __name__ == "__main__":
    main()
